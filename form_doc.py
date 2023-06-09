from pathlib import Path
from subprocess import Popen
from tkinter import filedialog as fd, messagebox as mb
import docx
import json


def form_doc(
        addressees, npa_type, npa_title, npa_date, npa_num, publ_date,
        app_sheets
):
    NPA_DATA = {
        'ФКЗ': {
            'title_text': f'О Федеральном конституционном законе от {npa_date} \
№ {npa_num}',
            'case1': 'опубликован', 'case2': 'указанного Федерального \
конституционного закона',
            'long_type': 'Федеральный конституционный закон',
            'short_type': 'ФКЗ'
        },
        'ФЗ': {
            'title_text': f'О Федеральном законе от {npa_date} № {npa_num}',
            'case1': 'опубликован', 'case2': 'указанного Федерального закона',
            'long_type': 'Федеральный закон',
            'short_type': 'ФЗ'
        },
        'Указ': {
            'title_text': f'''Об Указе Президента Российской Федерации от \
{npa_date} № {npa_num}''',
            'case1': 'опубликован', 'case2': 'Указа',
            'long_type': 'Указ Президента Российской Федерации',
            'short_type': 'Указ'
        },
        'ППРФ': {
            'title_text': f'''О постановлении Правительства Российской \
Федерации от {npa_date} № {npa_num}''', 'case1': 'опубликовано',
            'case2': 'указанного постановления',
            'long_type': 'постановление Правительства Российской Федерации',
            'short_type': 'ППРФ'
        },
        'РПРФ': {
            'title_text': f'''О распоряжении Правительства Российской \
Федерации от {npa_date} № {npa_num}''', 'case1': 'опубликовано',
            'case2': 'указанного распоряжения',
            'long_type': 'распоряжение Правительства Российской Федерации',
            'short_type': 'РПРФ'
        },
        'Пост. КС РФ': {
            'title_text': f'''О Постановлении Конституционного Суда \
Российской Федерации от {npa_date} № {npa_num}''', 'case1': 'опубликовано',
            'case2': 'указанного Постановления',
            'long_type':
            'Постановление Конституционного Суда Российской Федерации',
            'short_type': 'Пост. КС РФ'
        },
        'Опр. КС РФ': {
            'title_text': f'''Об Определении Конституционного Суда Российской \
Федерации от {npa_date} № {npa_num}''', 'case1': 'опубликовано',
            'case2': 'указанного Определения',
            'long_type':
            'Определение Конституционного Суда Российской Федерации',
            'short_type': 'Опр. КС РФ'
        },
        'Приказ Минстроя': {
            'title_text':
            f'О приказе Минстроя России от {npa_date} № {npa_num}',
            'case1': 'опубликован', 'case2': 'приказа',
            'long_type': 'приказ Министерства строительства и жилищно-\
коммунального хозяйства Российской Федерации',
            'short_type': 'приказ Минстроя'
        },
        'Приказ Минфина': {
            'title_text':
            f'О приказе Минфина России от {npa_date} № {npa_num}',
            'case1': 'опубликован', 'case2': 'приказа',
            'long_type': 'приказ Министерства финансов Российской Федерации',
            'short_type': 'приказ Минфина'
        }
    }
    current_npa_type = NPA_DATA[npa_type]
    doc_path = Path(__file__).with_name('doc_template.docx')
    doc = docx.Document(doc_path)
    table = doc.tables[0]
    address = table.cell(0, 2)
    title = table.cell(1, 0)
    exclamation = doc.paragraphs[1]
    first_para = doc.paragraphs[3]
    second_para = doc.paragraphs[4]
    appendix = doc.paragraphs[5]
    address.paragraphs[0].text = addressees[0].position
    address.paragraphs[1].text = addressees[0].name_with_initials
    for i in range(1, len(addressees)):
        address.add_paragraph(addressees[i].position)
        address.add_paragraph(addressees[i].name_with_initials)
    for para in address.paragraphs:
        para.style = 'addr_style'
    title.paragraphs[1].text = current_npa_type['title_text']
    title.paragraphs[1].style = 'title_style'
    if len(addressees) == 1:
        exclamation.text = addressees[0].address
    else:
        exclamation.text = 'Уважаемые коллеги!'
    first_para.text = f'Сообщаю, что {publ_date} на официальном \
Интернет-портале правовой информации pravo.gov.ru {current_npa_type["case1"]} \
{npa_title}.'
    second_para.text = f'Направляю копию {current_npa_type["case2"]} для \
сведения и учета в работе.'
    appendix.text = f'Приложение: на {app_sheets} л. в 1 экз.'
    if npa_type == 'Приказ Минстроя':
        savename = (
            f'{", ".join([elem.name_with_initials for elem in addressees])} \
{current_npa_type["short_type"]} {npa_num[:-3]}_пр.docx'
        )
    else:
        savename = (
            f'{", ".join([elem.name_with_initials for elem in addressees])} \
{current_npa_type["short_type"]} {npa_num}.docx'
        )
    config_path = Path(__file__).with_name('config.json')
    with config_path.open('r+') as config:
        data = json.load(config)
        existing_path = data.get('generated_docs_folder')
        if existing_path is not None:
            savepath = Path(existing_path)
        else:
            mb.showinfo(
                'Выбор папки', 'Выберите расположение папки для сохранения \
файлов служебных записок'
            )
            savepath = fd.askdirectory()
            if savepath == '':
                return mb.showerror(
                    'Выбор папки', 'Папка не выбрана, сохранение документа \
невозможно')
            data['generated_docs_folder'] = savepath
            config.seek(0)
            json.dump(data, config)
            config.truncate()
    doc.save(Path(savepath, savename))
    Popen([Path(savepath, savename)], shell=True)


if __name__ == '__main__':
    doc = docx.Document(Path('doc_template.docx'))
    exclamation = doc.paragraphs[1]
    print(exclamation.text)
